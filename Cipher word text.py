import sys
import docx

#select to "Encrypt" or "Decrypt"
Type = "Encrypt"
document = docx.Document(r"C:\Users\Desktop\Test.docx") #Select directory of word document
#Encryption Key, must have same length as charlist, and should have one of every character
Key = 'a9z8b7c6d 5e4f3g2h1i0j/k*l-m=y.x,w+vAuBtCsDrnEoFpGqHZIYJXKWLMUVNOPSTQR'
Charlist = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789/*-+=., '

#extract document text into string
fullText = []
for para in document.paragraphs:
    fullText.append(para.text)
    Input = ('\n'.join(fullText))

#Encryption
if Type == "Encrypt" :
  document._body.clear_content()
  trans = str.maketrans(Key, Charlist)
  Encryptedstr = Input.translate(trans)
  document.add_paragraph(Encryptedstr)
  print (Encryptedstr)

#Decrption
if Type == "Decrypt" :
  document._body.clear_content()
  trans = str.maketrans(Charlist, Key)
  Decryptedstr = Input.translate(trans)
  document.add_paragraph(Decryptedstr)
  print (Decryptedstr)

#Output save document
document.save(r"C:\Users\Desktop\Test1.docx") #Select directory of word document output
